"""Export writing runs to DOCX and PDF formats."""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from app.core.logging import get_logger

logger = get_logger(__name__)


_STYLE_CONFIGS = {
    "paper": {"font": "Times New Roman", "size": 12, "line_spacing": 2.0, "margins_cm": 2.5},
    "article": {"font": "Calibri", "size": 11, "line_spacing": 1.5, "margins_cm": 2.5},
    "summary": {"font": "Calibri", "size": 11, "line_spacing": 1.2, "margins_cm": 2.0},
    "draft": {"font": "Calibri", "size": 11, "line_spacing": 1.2, "margins_cm": 2.0},
}


def _get_style(doc_type: str) -> dict:
    return _STYLE_CONFIGS.get(doc_type, _STYLE_CONFIGS["article"])


def export_docx(markdown_content: str, doc_type: str = "article") -> Path:
    """Convert markdown content to a .docx file. Returns path to temp file."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
        from docx.enum.text import WD_LINE_SPACING
    except ImportError as e:
        raise RuntimeError("python-docx not installed. Add python-docx to backend dependencies.") from e

    cfg = _get_style(doc_type)
    doc = Document()

    # Set margins
    for section in doc.sections:
        margin = Cm(cfg["margins_cm"])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    lines = markdown_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            _set_font(p, cfg["font"], cfg["size"] + 4)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            _set_font(p, cfg["font"], cfg["size"] + 2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            _set_font(p, cfg["font"], cfg["size"] + 1)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles["Body Text"]
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            _apply_inline(p, cfg)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            _set_para_font(p, cfg["font"], cfg["size"])
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            p = doc.add_paragraph(text, style="List Number")
            _set_para_font(p, cfg["font"], cfg["size"])
        elif line.strip() == "":
            pass  # blank line — paragraph break handled by next content line
        else:
            p = doc.add_paragraph()
            _set_para_font(p, cfg["font"], cfg["size"])
            _apply_inline_runs(p, line, cfg)
            if cfg["line_spacing"] == 2.0:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            elif cfg["line_spacing"] == 1.5:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        i += 1

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    logger.info("DOCX export complete", path=tmp.name)
    return Path(tmp.name)


def export_pdf(markdown_content: str, doc_type: str = "article") -> Path:
    """Convert markdown content to a .pdf file. Returns path to temp file."""
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    except ImportError as e:
        raise RuntimeError("reportlab not installed. Add reportlab to backend dependencies.") from e

    cfg = _get_style(doc_type)
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    margin = cfg["margins_cm"] * cm

    doc = SimpleDocTemplate(
        tmp.name,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )

    styles = getSampleStyleSheet()
    font_name = "Times-Roman" if "Times" in cfg["font"] else "Helvetica"
    leading = cfg["size"] * cfg["line_spacing"] * 1.2

    style_normal = ParagraphStyle(
        "CustomNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=cfg["size"],
        leading=leading,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    style_h1 = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontName=font_name + "-Bold" if "Times" in cfg["font"] else "Helvetica-Bold",
        fontSize=cfg["size"] + 4,
        spaceAfter=8,
        spaceBefore=16,
        alignment=TA_CENTER,
    )
    style_h2 = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontName=font_name + "-Bold" if "Times" in cfg["font"] else "Helvetica-Bold",
        fontSize=cfg["size"] + 2,
        spaceAfter=6,
        spaceBefore=12,
    )
    style_h3 = ParagraphStyle(
        "CustomH3",
        parent=styles["Heading3"],
        fontName=font_name,
        fontSize=cfg["size"] + 1,
        spaceAfter=4,
        spaceBefore=8,
    )
    style_quote = ParagraphStyle(
        "CustomQuote",
        parent=style_normal,
        leftIndent=cm * 1.0,
        rightIndent=cm * 1.0,
        fontName=font_name + "-Oblique" if "Times" in cfg["font"] else "Helvetica-Oblique",
        textColor=colors.HexColor("#4b5563"),
    )
    style_bullet = ParagraphStyle(
        "CustomBullet",
        parent=style_normal,
        leftIndent=cm * 0.5,
        bulletIndent=0,
    )

    story = []
    for line in markdown_content.split("\n"):
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        if line.startswith("# "):
            story.append(Paragraph(_rl_escape(line[2:].strip()), style_h1))
        elif line.startswith("## "):
            story.append(Paragraph(_rl_escape(line[3:].strip()), style_h2))
        elif line.startswith("### "):
            story.append(Paragraph(_rl_escape(line[4:].strip()), style_h3))
        elif line.startswith("> "):
            story.append(Paragraph(_rl_escape(line[2:].strip()), style_quote))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph("• " + _rl_escape(line[2:].strip()), style_bullet))
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^(\d+)\. ", r"\1. ", line)
            story.append(Paragraph(_rl_escape(text), style_normal))
        else:
            story.append(Paragraph(_rl_format_inline(line), style_normal))

    doc.build(story)
    logger.info("PDF export complete", path=tmp.name)
    return Path(tmp.name)


def _rl_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _rl_format_inline(text: str) -> str:
    text = _rl_escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def _set_font(para: object, font_name: str, size: int) -> None:
    try:
        from docx.shared import Pt
        for run in para.runs:  # type: ignore[attr-defined]
            run.font.name = font_name
            run.font.size = Pt(size)
    except Exception:
        pass


def _set_para_font(para: object, font_name: str, size: int) -> None:
    try:
        from docx.shared import Pt
        para.style.font.name = font_name  # type: ignore[attr-defined]
        for run in para.runs:  # type: ignore[attr-defined]
            run.font.name = font_name
            run.font.size = Pt(size)
    except Exception:
        pass


def _apply_inline(para: object, cfg: dict) -> None:
    try:
        from docx.shared import Pt
        for run in para.runs:  # type: ignore[attr-defined]
            run.font.name = cfg["font"]
            run.font.size = Pt(cfg["size"])
    except Exception:
        pass


def _apply_inline_runs(para: object, line: str, cfg: dict) -> None:
    """Parse **bold** and *italic* markers into docx runs."""
    try:
        from docx.shared import Pt
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
        for match in pattern.finditer(line):
            if match.group(2):
                run = para.add_run(match.group(2))  # type: ignore[attr-defined]
                run.bold = True
            elif match.group(3):
                run = para.add_run(match.group(3))  # type: ignore[attr-defined]
                run.italic = True
            else:
                run = para.add_run(match.group(4) or "")  # type: ignore[attr-defined]
            run.font.name = cfg["font"]
            run.font.size = Pt(cfg["size"])
    except Exception:
        para.add_run(line)  # type: ignore[attr-defined]
